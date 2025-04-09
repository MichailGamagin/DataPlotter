#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import sys
import os
import yaml

from pathlib import Path
from PyQt5.QtWidgets import (
    QApplication,
    QMainWindow,
    QWidget,
    QVBoxLayout,
    QHBoxLayout,
    QPushButton,
    QFileDialog,
    QStackedWidget,
    QToolBar,
    QLabel,
    QMessageBox,
    QSizePolicy,
    QShortcut

)
from PyQt5.QtGui import QIcon, QKeySequence
from PyQt5.QtCore import Qt

from src.gui.styles import MY_LINE_EDIT_STYLE, STACK_WIDGET_STYLE
from src.core.constants import BASE_DIR, ENCODING, DEFAULT_FILE_PATH, SAVE_FILE, DEFAULT_DIR, ICONS_DIR

from src.core.data_loader import load_data_from
from src.gui.views.components.lyne_edit import MyLineEdit
from src.gui.views.components.left_panel import LeftPanel
from src.gui.views.components.plot_area import PlotArea
from src.gui.views.dialogs.message import MessageWindow
from src.gui.views.dialogs.progress import MyProgressDialog
from src.gui.views.word.settings import WordSettings
from src.utils.logger import Logger

logger = Logger.get_logger(__name__)

class MainWindow(QMainWindow):
    """
    Главное окно приложения для построения графиков.

    Attributes:
        pages (list): Список словарей, содержащих информацию о каждой странице (графике).
        current_page (int): Индекс текущей активной страницы.
        data_file_path (str): Путь к файлу с данными.
        data (pd.DataFrame): DataFrame с данными для построения графиков.
        stack (QStackedWidget): Виджет для отображения страниц.
        new_btn (QPushButton): Кнопка для добавления новой страницы.
        prev_btn (QPushButton): Кнопка для перехода на предыдущую страницу.
        next_btn (QPushButton): Кнопка для перехода на следующую страницу.
        save_all_btn (QPushButton): Кнопка для сохранения всех графиков.
        save_state_btn (QPushButton): Кнопка для сохранения состояния приложения.
        load_state_btn (QPushButton): Кнопка для загрузки состояния приложения.
        word_btn (QPushButton): Кнопка для экспорта графиков в Word.
        path_lbl (QLabel): Метка для отображения текста "Путь к файлу:".
        path_ent (MyLineEdit): Поле ввода для указания пути к файлу с данными.
    """

    def __init__(self):
        super().__init__()
        logger.info("Инициализация MainWindow")
        self.pages = []
        self.params = {}
        self.alternative_captions = {}
        self.setFixedSize(1350, 750)
        self.current_page = 0
        self.data_file_path = DEFAULT_FILE_PATH
        self.data = load_data_from(self.data_file_path, ENCODING)
        self.init_ui()
        self._load_state()
        logger.info("Интерфейс MainWindow успешно инициализирован")

    def init_ui(self):
        logger.info(f"Инициализация пользовательского интерфейса MainWindow")
        self.stack = QStackedWidget()
        self.setCentralWidget(self.stack)
        self.setStyleSheet(STACK_WIDGET_STYLE)
        # Toolbar
        toolbar = QToolBar()
        self.addToolBar(toolbar)

        # Navigation buttons
        self.new_btn = QPushButton("")
        self.new_btn.setIcon(QIcon(os.path.join(ICONS_DIR, "icons", "add_graph.png")))
        self.new_btn.setToolTip("Добавить график")
        self.new_btn.clicked.connect(self.add_page)

        self.prev_btn = QPushButton()
        self.prev_btn.setIcon(QIcon(os.path.join(ICONS_DIR, "icons", "prev.png")))
        self.prev_btn.setToolTip("Предыдущий график")
        self.prev_btn.clicked.connect(self.prev_page)
        self.shortcut_left = QShortcut(QKeySequence('Left'), self)
        self.shortcut_left.activated.connect(self.prev_page)

        self.next_btn = QPushButton()
        self.next_btn.setIcon(QIcon(os.path.join(ICONS_DIR, "icons", "next.png")))
        self.next_btn.setToolTip("Следующий график")
        self.next_btn.clicked.connect(self.next_page)
        self.shortcut_right = QShortcut(QKeySequence("Right"), self)
        self.shortcut_right.activated.connect(self.next_page)

        self.save_all_btn = QPushButton("")
        self.save_all_btn.clicked.connect(self.save_all)
        self.save_all_btn.setIcon(
            QIcon(os.path.join(ICONS_DIR, "icons", "save_all.png"))
        )
        self.save_all_btn.setToolTip("Сохранить все графики")

        self.save_state_btn = QPushButton("")
        self.save_state_btn.clicked.connect(self.save_state)
        self.save_state_btn.setIcon(
            QIcon(os.path.join(ICONS_DIR, "icons", "save_state.png"))
        )
        self.save_state_btn.setToolTip("Сохранить состояние")

        self.load_state_btn = QPushButton("")
        self.load_state_btn.clicked.connect(self.load_state)
        self.load_state_btn.setIcon(
            QIcon(os.path.join(ICONS_DIR, "icons", "load_state.png"))
        )
        self.load_state_btn.setToolTip("Загрузить состояние")

        self.word_btn = QPushButton("")
        self.word_btn.clicked.connect(self.export_to_word)
        self.word_btn.setIcon(QIcon(os.path.join(ICONS_DIR, "icons", "word.png")))
        self.word_btn.setToolTip("Перенести графики в Word")
        self.word_settings_btn = QPushButton("")
        self.word_settings_btn.clicked.connect(self.open_settings)
        self.word_settings_btn.setIcon(
            QIcon(os.path.join(ICONS_DIR, "icons", "settings.jpg"))
        )
        self.word_settings_btn.setToolTip("Настройки Word")

        spacer = QWidget()
        spacer.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)

        self.path_lbl = QLabel("Путь к файлу:")
        self.path_lbl.setStyleSheet(
            "color: black;\n" "font-size: 14px;\n" "font-weight: 700;"
        )
        self.path_ent = MyLineEdit()
        self.path_ent.setPlaceholderText(r"C:\Users\...\res_cyclic.txt")
        self.path_ent.setMaximumWidth(600)
        self.path_ent.setAcceptDrops(True)
        self.path_ent.setStyleSheet(MY_LINE_EDIT_STYLE)
        self.path_ent.editingFinished.connect(self.load_data)

        toolbar.addWidget(self.new_btn)
        toolbar.addSeparator()
        toolbar.addWidget(self.prev_btn)
        toolbar.addWidget(self.next_btn)
        toolbar.addSeparator()
        toolbar.addWidget(self.save_state_btn)
        toolbar.addWidget(self.load_state_btn)
        toolbar.addSeparator()
        toolbar.addWidget(self.save_all_btn)
        toolbar.addSeparator()
        toolbar.addWidget(self.word_btn)
        toolbar.addWidget(self.word_settings_btn)
        toolbar.addWidget(spacer)
        toolbar.addWidget(self.path_lbl)
        toolbar.addWidget(self.path_ent)

        if self.data_file_path == "":
            self.add_page()
        self.update_buttons()

        # Window settings
        self.setWindowTitle("Data Plotter")
        # self.setGeometry(300, 300, 1300, 700)
        self.setWindowIcon(QIcon(os.path.join(ICONS_DIR, "icons", "main_icon.png")))
        logger.info("Пользовательский интерфейс MainWindow успешно инициализирован")
        self.show()

    def load_data(self):
        """Загружает данные из файла, указанного в self.path_ent"""
        file_path = self.path_ent.text()
        logger.info(f"Попытка загрузки данных из файла: {file_path}")
        try:
            if file_path.endswith(".yaml"):
                with open(file_path, "r", encoding=ENCODING) as f:
                    state = yaml.load(f, Loader=yaml.FullLoader)
                    file_path = state.get("data_file_path", DEFAULT_FILE_PATH)
                    self.path_ent.setText(file_path)

            self.data = load_data_from(file_path, ENCODING)
            for page in self.pages:
                for combo in page["left"].combos:
                    combo.clear()
                    combo.blockSignals(True)
                    combo.addItems(self.data.columns[1:])
                    combo.setCurrentIndex(-1)
                    combo.blockSignals(False)
                page["right"].data = self.data
                page["right"].clear_graph()

                self.update_graph()
        except FileNotFoundError:
            logger.error(f"Файл данных не найден: {file_path}")
            QMessageBox.critical(self, "Ошибка", f"Файл данных не найден: {file_path}")
        except Exception as e:
            logger.error(f"Ошибка при загрузке данных из файла: {file_path}. Текст ошибки: {str(e)}", exc_info=True)
            QMessageBox.critical(self, "Ошибка", f"Ошибка при загрузке данных из файла: {file_path}\nТекст ошибки: {str(e)}")

    def export_to_word(self):
        """Экспорт всех графиков в документ Word"""
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.enum.style import WD_STYLE_TYPE
        # Запрос пути сохранения
        options = QFileDialog.Options()
        # options |= QFileDialog.DontUseNativeDialog
        path = self.path_ent.text()
        if path:
            home_dir = Path(path).parent.as_posix()
        else:
            home_dir = os.path.expanduser(DEFAULT_DIR)
        file_name, _ = QFileDialog.getSaveFileName(
            self,
            "Сохранить документ Word",
            home_dir,
            "Word Files (*.docx);;All Files (*)",
            options=options,
        )

        if not file_name:
            return
        logger.info(f"Сохранение графиков в Word в директорию: {file_name}")

        try:
            # Сохраняем текущую страницу
            original_page = self.current_page
            # Махинации с Word
            doc = Document()
            section = doc.sections[0]
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)
            # Основной стиль документа
            style = doc.styles["Normal"]
            style.font.name = self.params.get("font", "Times New Roman")
            font_size = self.params.get("font-size", "11")
            int_after = self.params.get("int-after", "0")
            int_before = self.params.get("int-before", "0")
            pic_width = self.params.get("pic-width", "16.0")
            pic_height = self.params.get("pic-height", "9.5")
            line_spacing = self.params.get("line-spacing", 'Одинарный')
            if isinstance(font_size, (str, float, int)):
                try:
                    font_size = Pt(float(font_size))
                except ValueError:
                    font_size = Pt(11)
                    msg = MessageWindow(
                        f"Ошибка в размере шрифта",
                    )
                    msg.exec_()
            if isinstance(int_after, (int, float)):
                try:
                    int_after = Pt(float(int_after))
                except ValueError:
                    int_after = Pt(0)
                    msg = MessageWindow(
                        f"Ошибка в интервале после",
                    )
                    msg.exec_()
            if isinstance(int_before, (int, float)):
                try:
                    int_before = Pt(float(int_before))
                except ValueError:
                    int_before = Pt(0)
                    msg = MessageWindow(
                        f"Ошибка в интервале до",
                    )
                    msg.exec_()
            if isinstance(pic_width, (str, int, float)):
                try:
                    pic_width = Cm(float(pic_width))
                except ValueError:
                    pic_width = Cm(16.0)
                    msg = MessageWindow(
                        f"Ошибка в ширине рисунка",
                        'Установлена ширина 16.0 см',
                    )
                    msg.exec_()
            if isinstance(pic_height, (str, int, float)):
                try:
                    pic_height = Cm(float(pic_height))
                except ValueError:
                    pic_height = Cm(9.5)
                    msg = MessageWindow(
                        f"Ошибка в высоте рисунка", 
                        "Установлена высота 9.5 см",
                    )
                    msg.exec_()
            style.font.size = font_size
            # Стиль списка линий под подписью рисунка
            obj_styles = doc.styles
            list_style = obj_styles.add_style("ListStyle", WD_STYLE_TYPE.PARAGRAPH)
            list_style.base_style = obj_styles["Normal"]
            list_style.font.name = self.params.get("font", "Times New Roman")
            list_style.font.size =  font_size
            # Добавляем изменение межстрочного интервала и интервалов до и после
            paragraph_format = list_style.paragraph_format
            if line_spacing == "Одинарный":
                paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            else:
                paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            paragraph_format.space_after = int_after
            paragraph_format.space_before = int_before
            # Временная папка для изображений
            temp_dir = Path("./temp_grf")
            temp_dir.mkdir(exist_ok=True, parents=True)

            for idx, page in enumerate(self.pages):
                # Переключаемся на страницу
                self.current_page = idx
                self.stack.setCurrentIndex(idx)
                QApplication.processEvents()

                # Обновляем график
                self.update_graph()
                page["right"].canvas.draw_idle()
                QApplication.processEvents()

                # Сохраняем временное изображение
                img_path = temp_dir / f"graph_{idx}.png"
                page["right"].canvas.figure.savefig(img_path, dpi=300)

                # Добавляем в документ
                # 1. Изображение
                para_img = doc.add_paragraph()
                para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para_img.add_run()
                run.add_picture(str(img_path), width=pic_width, height=pic_height)

                # Получаем подписи линий
                graph_header = f"График №{idx + 1}"
                alternative_caption = self.alternative_captions.get(graph_header)

                list_item = doc.add_paragraph()
                list_item.alignment = WD_ALIGN_PARAGRAPH.CENTER
                list_item.style = list_style

                if alternative_caption:
                    list_item.add_run(alternative_caption)
                else:
                    labels = [
                        combo.currentText().split(",")[0].strip()
                        for combo in page["left"].combos
                        if combo.currentText()
                    ]
                    for i, label in enumerate(labels, 1):
                        list_item.add_run(f"{i} - {label}")
                        if i < len(labels):
                            list_item.add_run("\n")

                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_text = f'{self.params["pict"]} {self.params["num-section"]}.{int(self.params["first-pic"])+idx} - {self.params["mode-name"]}'
                caption.add_run(caption_text).bold = False
                caption.style = list_style

            # Удаляем временные файлы
            for f in temp_dir.glob("*.png"):
                f.unlink()
            temp_dir.rmdir()

            # Сохраняем документ
            if not file_name.endswith('.docx'):
                file_name = file_name + '.docx'
            doc.save(file_name)
            QMessageBox.information(
                self, "Успех", f"Графики успешно экспортированы в Word в директорию:\n{file_name}"
            )
            logger.info(
                f"Графики успешно экспортированы в Word в директорию: {file_name}"
            )
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка экспорта: {str(e)}")

        finally:
            # Восстанавливаем состояние
            self.current_page = original_page
            self.stack.setCurrentIndex(original_page)
            self.update_buttons()

    def open_settings(self):
        self.word_settings = WordSettings(parent=self)
        self.word_settings.setWindowModality(Qt.ApplicationModal)
        self.word_settings.show()

    def add_page(self):
        logger.info(f"Добавление страницы")
        page = QWidget()
        layout = QHBoxLayout(page)
        vlayout = QVBoxLayout()

        left_panel = LeftPanel(self.data, self)
        plot_area = PlotArea(self.data, self)

        vlayout.addWidget(left_panel, 1)
        vlayout.addStretch(10)
        layout.addLayout(vlayout, 1)
        layout.addWidget(plot_area, 4)

        self.pages.append({"widget": page, "left": left_panel, "right": plot_area})
        self.stack.addWidget(page)
        self.current_page = len(self.pages) - 1
        self.stack.setCurrentIndex(self.current_page)
        self.pages[self.current_page]["left"].update_label()
        self.update_buttons()
        self.stack.setCurrentIndex(self.current_page)
        logger.info(f"Добавлена страница №{self.current_page + 1}")

    def update_buttons(self):
        self.prev_btn.setEnabled(self.current_page > 0)
        self.next_btn.setEnabled(self.current_page < len(self.pages) - 1)

    def update_graph(self):
        combos = self.pages[self.current_page]["left"].combos
        for i in range(len(combos)):
            self.plot_selection(i)

    def prev_page(self):
        if self.current_page > 0:
            self.current_page -= 1
            self.stack.setCurrentIndex(self.current_page)
            self.pages[self.current_page]["left"].update_label()
            self.update_buttons()
            self.update_graph()

    def next_page(self):
        if self.current_page < len(self.pages) - 1:
            self.current_page += 1
            self.stack.setCurrentIndex(self.current_page)
            self.pages[self.current_page]["left"].update_label()
            self.update_buttons()
            self.update_graph()

    def plot_selection(self, combo_idx: int):
        current_page = self.pages[self.current_page]
        combo = current_page["left"].combos[combo_idx]
        selected_col = combo.currentText()
        if selected_col:
            current_page["right"].update_plot(combo_idx, selected_col)

    def save_state(self):
        options = QFileDialog.Options()
        # options |= QFileDialog.DontUseNativeDialog
        path = self.path_ent.text()
        if path:
            home_dir = Path(path).parent.as_posix()
        else:
            home_dir = os.path.expanduser(DEFAULT_DIR)
        file_name, _ = QFileDialog.getSaveFileName(
            self,
            "Сохранить состояние",
            home_dir,
            "All Files (*)",
            options=options,
        )

        if not file_name:
            return
        if not file_name.endswith(".yaml"):
            file_name = file_name + ".yaml"
        logger.info(f"Сохранение состояния в директорию: {file_name}")
        if file_name:
            path = Path(self.path_ent.text())
            state = {
                "data_file_path": path.as_posix(),
                "_Word": self.params,
                "pages": [],
            }
            for i, page_data in enumerate(self.pages):
                # Получаем заголовок графика
                graph_header = f"График №{i+1}"
                # Получаем альтернативную подпись из self.alternative_captions
                alternative_caption = self.alternative_captions.get(graph_header, "")
                page_state = {
                    "Number_graph": graph_header,
                    "Symbol": [
                        str(page_data["right"].group.currentText()),
                        str(page_data["right"].sizing_cmb.currentText()),
                    ],
                    "Axis_settings": {
                        "X": str(page_data["right"].x_settings.currentText()),
                        "Y": str(page_data["right"].y_settings.text()),
                        "Frequency": str(page_data["right"].markers.currentText()),
                        "X_grid_lines": str(page_data["right"].x_spacing_grid_spinBox.value()),
                    },
                    "Lists": [],
                    "alternative_caption": alternative_caption,
                }
                for combo in page_data["left"].combos:
                    text = combo.currentText()
                    page_state["Lists"].append(text)
                state["pages"].append(page_state)

            with open(file_name, "w", encoding="cp1251") as f:
                yaml.dump(state, f, indent=2, allow_unicode=True)
            logger.info(f"Состояние успешно сохранено в директорию {file_name}")

    def _load_state(self, file_path=SAVE_FILE):
        logger.info("Загрузка состояния...")
        try:
            logger.debug("Загрузка состояния...")
            progress = MyProgressDialog(title="Загрузка состояния", parent=self)
            progress.show()
            QApplication.processEvents()

            with open(file_path, "r", encoding=ENCODING) as f:
                state = yaml.load(f, Loader=yaml.FullLoader)

            # 1. Очистка текущего состояния
            while self.pages:
                page = self.pages.pop()
                self.stack.removeWidget(page["widget"])
                QApplication.processEvents()

            # 2. Загрузка данных
            self.path_ent.blockSignals(True)
            self.data_file_path = state.get("data_file_path", DEFAULT_FILE_PATH)
            self.params = state.get("_Word", {})
            self.path_ent.setText(self.data_file_path)
            self.path_ent.blockSignals(False)
            self.data = load_data_from(self.data_file_path, ENCODING)
            self.alternative_captions = {}

            # 3. Воссоздание страниц
            total_pages = len(state.get("pages", []))
            progress.setMaximum(total_pages)
            for i, page_state in enumerate(state.get("pages", [])):
                progress.setValue(i)
                progress.setLabelText(f"Загрузка страницы {i}/{total_pages}")
                QApplication.processEvents()

                if progress.wasCanceled():
                    break

                self.current_page = 0
                self.add_page()
                current_page = self.pages[-1]

                # Загрузка параметров правой части
                current_page["right"].x_settings.setCurrentText(
                    page_state["Axis_settings"]["X"]
                )
                current_page["right"].y_settings.setText(
                    page_state["Axis_settings"]["Y"]
                )
                current_page["right"].markers.setCurrentText(
                    page_state["Axis_settings"]["Frequency"]
                )
                current_page["right"].x_spacing_grid_spinBox.setValue(int(
                    page_state["Axis_settings"]["X_grid_lines"]
                )
                )
                current_page["right"].group.setCurrentText(page_state["Symbol"][0])
                current_page["right"].sizing_cmb.setCurrentText(page_state["Symbol"][1])

                # Загрузка параметров левой части
                for j, combo_text in enumerate(page_state["Lists"]):
                    if j < len(current_page["left"].combos):
                        current_page["left"].combos[j].setCurrentText(combo_text)

                # Загрузка альтернативной подписи
                graph_header = f"{page_state['Number_graph']}"
                self.alternative_captions[graph_header] = page_state.get("alternative_caption", "")
                self.update_graph()
                QApplication.processEvents()

            # 4. Восстановление позиции
            if self.pages:
                self.current_page = 0
                self.stack.setCurrentIndex(0)
                self.update_buttons()
                self.pages[0]["left"].update_label()
            logger.info("Состояние успешно загружено")

        except Exception as e:
            if isinstance(e, FileNotFoundError):
                return
            else:
                QMessageBox.critical(self, "Ошибка", f"Ошибка загрузки: {str(e)}")
            logger.error(f"Ошибка при загрузке состояния: {str(e)}", exc_info=True)
        finally:
            progress.close()

    def load_state(self):
        options = QFileDialog.Options()
        options |= QFileDialog.ReadOnly
        # options |= QFileDialog.DontUseNativeDialog

        path = self.path_ent.text()
        if path:
            home_dir = Path(path).parent.as_posix()
        else:
            home_dir = os.path.expanduser(DEFAULT_DIR)

        self.state_file_path, _ = QFileDialog.getOpenFileName(
            self, "Загрузить состояние", home_dir, "All Files (*)", options=options
        )
        if self.state_file_path != "":
            try:
                self._load_state(self.state_file_path)

            except Exception as e:
                msg = MessageWindow(
                    f"Ошибка загрузки состояния из файла:\n{self.state_file_path}",
                    f"Текст ошибки: {str(e)}",
                )
                msg.exec_()

    def save_all(self):
        """Сохраняет все графики в выбранную пользователем директорию."""
        # Запрос директории сохранения
        options = QFileDialog.Options()
        options |= QFileDialog.ShowDirsOnly  # Показываем только папки
        # options |= QFileDialog.DontUseNativeDialog
        path = self.path_ent.text()
        if path:
            home_dir = Path(path).parent.as_posix()
        else:
            home_dir = os.path.expanduser(DEFAULT_DIR)
        directory = QFileDialog.getExistingDirectory(
            self, "Выберите папку для сохранения графиков", home_dir, options=options
        )

        if not directory:
            return  # Пользователь отменил выбор

        logger.info(f"Сохранение всех графиков в директорию: {directory}")
        # Сохраняем текущую активную страницу
        original_page = self.current_page
        try:
            # Создаем папку если нужно
            for idx, page in enumerate(self.pages):
                self.current_page = idx
                self.stack.setCurrentIndex(idx)

                # Обновляем интерфейс страницы
                QApplication.processEvents()

                # Принудительно обновляем графики
                self.update_graph()
                page["left"].update_label()
                page["right"].canvas.draw_idle()
                QApplication.processEvents()

                number = (
                    page["left"].num_page.text().replace(" ", "_").replace("/", "-")
                )
                filename = os.path.join(directory, f"{number}.png")

                page["right"].canvas.figure.savefig(
                    filename,
                    format="png",
                    dpi=300,
                )
            QMessageBox.information(
                self, "Успешно", f"Графики сохранены в:\n{directory}"
            )
            self.update_buttons()
            self.current_page = original_page
            self.stack.setCurrentIndex(original_page)
            logger.info(f"Успешное сохранение всех графиков в директорию: {directory}")

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка сохранения: {str(e)}")
            logger.error(f"Ошибка при сохранении всех графиков: {str(e)}", exc_info=True)
        finally:
            self.current_page = original_page
            self.stack.setCurrentIndex(original_page)
            self.update_buttons()


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    sys.exit(app.exec_())
