from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE

class Word:
    """Класс для экспортирования графиков в Word"""

    def __init__(self) -> None:
        self.doc = Document()

    def set_margin(self, top: Cm = Cm(1), bottom: Cm = Cm(1)):
        """Устанавливает верхний и нижний отступ на листе"""
        self.section = self.doc.sections[0]
        self.section.top_margin = top
        self.section.bottom_margin = bottom

    def set_style_document(self):
        """ "Устанавливает стиль документа"""
        self.style = self.doc.styles["Normal"]

    def get_style_list(self):
        return self.list_style

    def set_style_list(
        self,
        font: str = "Times New Roman",
        font_size: float | int | str = 10,
        line_spacing: str = "Одинарный",
        int_before: float | int | str = 0,
        int_after: float | int | str = 0,
    ):
        """Устанавливает стиль списка линий под рисунком"""
        if self.validation_params(
            **{
                "font_size": font_size,
                "int_before": int_before,
                "int_after": int_after,
            }
        ):
            self.list_style = self.doc.styles.add_style(
                "ListStyle", WD_STYLE_TYPE.PARAGRAPH
            )
            self.list_style.font.name = font
            self.list_style.font.size = Pt(float(font_size))
            paragraph_format = self.list_style.paragraph_format
            if line_spacing == "Одинарный":
                paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            else:
                paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            paragraph_format.space_after = Pt(float(int_after))
            paragraph_format.space_before = Pt(float(int_before))

    def add_image(
        self,
        image_path: Path = Path("./temp_grf"),
        pic_width: str | int | float = 16.0,
        pic_height: str | int | float = 9.0,
    ):
        """Добавление рисунка в документ"""

        if self.validation_params(**{"pic_width": pic_width, "pic_height": pic_height}):
            para_img = self.doc.add_paragraph()
            para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para_img.add_run()
            run.add_picture(
                image_path_or_stream=Path(image_path).as_posix(),
                width=Cm(float(pic_width)),
                height=Cm(float(pic_height)),
            )

    def set_caption(
        self,
        idx: int,
        labels: list[str],
        parameters: dict,
        alt_caption: str = "",
    ):
        """Устанавливает подписи к рисункам:\n
        Args:
        >>>    idx - индекс(номер рисунка)
        >>>    labels - список линий рисунка
        >>>    alt_caption - альтернативная подпись
        """
        list_item = self.doc.add_paragraph()
        list_item.alignment = WD_ALIGN_PARAGRAPH.CENTER
        list_item.style = self.list_style
        if alt_caption:
            list_item.add_run(alt_caption)
        else:
            for i, label in enumerate(labels, 1):
                list_item.add_run(f"{i} - {label}")
                if i < len(labels):
                    list_item.add_run("\n")

        caption = self.doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_text = f'{parameters["pict"]} {parameters["num-section"]}.{int(parameters["first-pic"])+idx} - {parameters["mode-name"]}'
        caption.add_run(caption_text).bold = False
        caption.style = self.get_style_list()

    def save_doc(self, path: Path):
        """Сохраняет документ по заданному пути"""
        if not path.endswith(".docx"):
            path = path + ".docx"
        self.doc.save(path)

    def validation_params(self, **kwargs) -> bool:
        """Метод валидации параметров"""
        return all(isinstance(value, (str, int, float)) for value in kwargs.values())


if __name__ == "__main__":
    path = Path("C:\\Users\\gamagin_MV\\Desktop\\DataPlotter_v12\\data\\samples\\word")
    word_document = Word(path.as_posix())
    par = {
        "font_size": 10,
        "int_before": "1",
        "int_after": "2",
        "pic_width": "16",
        "pic_height": "20",
    }
    print(word_document.set_style_list())
    print(word_document.get_style_list())
